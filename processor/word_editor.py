from docx import Document
from docx.shared import Cm, Inches
from docx.oxml.ns import qn
import os
from PIL import Image

def replace_text(doc: Document, replacements: dict, image_replacements: dict):
    """Reemplaza texto preservando estilos completamente."""
    def replace_in_paragraphs(paragraphs):
        for paragraph in paragraphs:
            full_text = ''.join(run.text for run in paragraph.runs)
            needs_replacement = any(key in full_text for key in replacements.keys())
            if needs_replacement:
                replace_text_preserve_formatting(paragraph, replacements)

    def replace_text_preserve_formatting(paragraph, replacements):
        char_formats = []
        full_text = ""
        
        for run in paragraph.runs:
            for char in run.text:
                char_formats.append({
                    'char': char,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'bold': run.font.bold,
                    'italic': run.font.italic,
                    'underline': run.font.underline,
                    'color': run.font.color.rgb if run.font.color.rgb else None,
                })
                full_text += char
        
        modified_text = full_text
        for key, value in replacements.items():
            modified_text = modified_text.replace(key, value)
        
        if modified_text != full_text:
            rebuild_paragraph_with_formatting(paragraph, char_formats, modified_text)

    def rebuild_paragraph_with_formatting(paragraph, char_formats, new_text):
        for run in paragraph.runs[:]:
            run._element.getparent().remove(run._element)
        
        if new_text and char_formats:
            base_format = char_formats[0]
            run = paragraph.add_run(new_text)
            apply_format_to_run(run, base_format)
        elif new_text:
            paragraph.add_run(new_text)

    def apply_format_to_run(run, format_info):
        if format_info.get('font_name'):
            run.font.name = format_info['font_name']
        if format_info.get('font_size'):
            run.font.size = format_info['font_size']
        if format_info.get('bold') is not None:
            run.font.bold = format_info['bold']
        if format_info.get('italic') is not None:
            run.font.italic = format_info['italic']
        if format_info.get('underline') is not None:
            run.font.underline = format_info['underline']
        if format_info.get('color'):
            run.font.color.rgb = format_info['color']

    def replace_in_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs)

    # Procesar todo el documento
    replace_in_paragraphs(doc.paragraphs)
    replace_in_tables(doc.tables)
    
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header is not None:
                try:
                    if hasattr(header, 'paragraphs') and header.paragraphs:
                        replace_in_paragraphs(header.paragraphs)
                    if hasattr(header, 'tables') and header.tables:
                        replace_in_tables(header.tables)
                except Exception as e:
                    print(f"⚠️ Error procesando header: {e}")
                    
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer is not None:
                try:
                    if hasattr(footer, 'paragraphs') and footer.paragraphs:
                        replace_in_paragraphs(footer.paragraphs)
                    if hasattr(footer, 'tables') and footer.tables:
                        replace_in_tables(footer.tables)
                except Exception as e:
                    print(f"⚠️ Error procesando footer: {e}")

def get_replacement_path(replacement_info):
    """Obtiene la ruta de reemplazo."""
    if isinstance(replacement_info, str):
        return replacement_info
    return replacement_info.get('path') if isinstance(replacement_info, dict) else None

def replace_placeholder_images(doc: Document, placeholder_replacements: dict):
    """Sistema de placeholders con reconstrucción completa de párrafos."""
    print("🔍 Buscando imágenes placeholder para reemplazar...")
    print(f"📋 Reemplazos disponibles: {list(placeholder_replacements.keys())}")
    
    if not placeholder_replacements:
        print("❌ No hay reemplazos de placeholders configurados")
        return 0
    
    # Verificar que al menos un archivo de reemplazo existe
    print("🔍 Verificando archivos de reemplazo...")
    valid_replacement_found = False
    for name, info in placeholder_replacements.items():
        path = get_replacement_path(info)
        if path and os.path.exists(path):
            print(f"✅ {name} -> {path}")
            valid_replacement_found = True
        else:
            print(f"❌ {name} -> {path} (NO EXISTE)")
    
    if not valid_replacement_found:
        print("❌ No hay archivos de reemplazo válidos")
        return 0
    
    total_replaced = 0
    
    # Ya no usamos replacement_queue - usamos directamente placeholder_replacements
    
    def process_all_paragraphs():
        nonlocal total_replaced
        
        # Procesar párrafos principales
        total_replaced += process_paragraph_list(doc.paragraphs, "contenido principal")
        
        # Procesar tablas
        for i, table in enumerate(doc.tables):
            for j, row in enumerate(table.rows):
                for k, cell in enumerate(row.cells):
                    total_replaced += process_paragraph_list(
                        cell.paragraphs, f"tabla {i+1}, celda ({j+1},{k+1})"
                    )
        
        # Procesar headers y footers
        for i, section in enumerate(doc.sections):
            headers = [
                (section.header, "header principal"),
                (section.first_page_header, "header primera página"), 
                (section.even_page_header, "header páginas pares")
            ]
            
            for header, name in headers:
                if header is not None:
                    try:
                        # Verificar que el header realmente existe y tiene contenido
                        if hasattr(header, '_element') and header._element is not None:
                            print(f"📋 Procesando {name} en sección {i+1}...")
                            
                            # Procesar párrafos del header
                            if hasattr(header, 'paragraphs') and header.paragraphs:
                                total_replaced += process_paragraph_list(
                                    header.paragraphs, f"{name} sección {i+1}"
                                )
                            
                            # Procesar tablas del header (CRÍTICO PARA TU CASO)
                            if hasattr(header, 'tables') and header.tables:
                                print(f"📊 Encontradas {len(header.tables)} tablas en {name}")
                                for j, table in enumerate(header.tables):
                                    print(f"   Procesando tabla {j+1} del header...")
                                    for k, row in enumerate(table.rows):
                                        for l, cell in enumerate(row.cells):
                                            cell_replaced = process_paragraph_list(
                                                cell.paragraphs, f"tabla {j+1} en {name} sección {i+1}, celda ({k+1},{l+1})"
                                            )
                                            if cell_replaced > 0:
                                                print(f"   ✅ Reemplazadas {cell_replaced} imágenes en tabla del header")
                                            total_replaced += cell_replaced
                    except Exception as e:
                        print(f"⚠️ Error procesando {name} en sección {i+1}: {e}")
            
            footers = [
                (section.footer, "footer principal"),
                (section.first_page_footer, "footer primera página"),
                (section.even_page_footer, "footer páginas pares")
            ]
            
            for footer, name in footers:
                if footer is not None:
                    try:
                        # Verificar que el footer realmente existe y tiene contenido
                        if hasattr(footer, '_element') and footer._element is not None:
                            print(f"📋 Procesando {name} en sección {i+1}...")
                            
                            # Procesar párrafos del footer
                            if hasattr(footer, 'paragraphs') and footer.paragraphs:
                                total_replaced += process_paragraph_list(
                                    footer.paragraphs, f"{name} sección {i+1}"
                                )
                            
                            # Procesar tablas del footer
                            if hasattr(footer, 'tables') and footer.tables:
                                for j, table in enumerate(footer.tables):
                                    for k, row in enumerate(table.rows):
                                        for l, cell in enumerate(row.cells):
                                            total_replaced += process_paragraph_list(
                                                cell.paragraphs, f"tabla en {name} sección {i+1}"
                                            )
                    except Exception as e:
                        print(f"⚠️ Error procesando {name} en sección {i+1}: {e}")
    
    def process_paragraph_list(paragraphs, location):
        replaced_count = 0
        for i, paragraph in enumerate(paragraphs):
            if paragraph_has_images(paragraph):
                if replace_images_in_paragraph(paragraph, f"{location}, párrafo {i+1}"):
                    replaced_count += 1
        return replaced_count
    
    def paragraph_has_images(paragraph):
        """Verifica si el párrafo contiene imágenes."""
        for run in paragraph.runs:
            for elem in run._element:
                if (elem.tag.endswith('drawing') or 
                    elem.tag.endswith('pict') or
                    'pic:pic' in elem.tag):
                    return True
        return False
    
    def replace_images_in_paragraph(paragraph, location):
        """Reemplaza imágenes en un párrafo reconstruyéndolo completamente."""
        try:
            # Verificar si hay reemplazos disponibles EN LA CONFIGURACIÓN ORIGINAL
            if not placeholder_replacements:
                return False
            # Buscar imágenes sin tocar el texto
            has_images = False
            for run in paragraph.runs:
                for elem in run._element:
                    if (elem.tag.endswith('drawing') or elem.tag.endswith('pict')):
                        has_images = True
                        break
                if has_images:
                    break
            
            if not has_images:
                return False
            
            # Reemplazar imagen sin reconstruir todo el párrafo
            placeholder_name, replacement_info = list(placeholder_replacements.items())[0]
            replacement_path = get_replacement_path(replacement_info)
            
            if replacement_path and os.path.exists(replacement_path):
                # Método directo: solo reemplazar la imagen
                for run in paragraph.runs:
                    for elem in run._element[:]:
                        if elem.tag.endswith('drawing') or elem.tag.endswith('pict'):
                            # Extraer dimensiones
                            original_dims = extract_image_info(elem)
                            new_dims = calculate_replacement_dimensions(
                                replacement_path, replacement_info, original_dims
                            )
                            
                            # Eliminar imagen vieja
                            elem.getparent().remove(elem)
                            
                            # Añadir imagen nueva SIN tocar el texto
                            width = Cm(new_dims['width_cm'])
                            height = Cm(new_dims['height_cm'])
                            run.add_picture(replacement_path, width=width, height=height)
                            
                            print(f"✅ Imagen reemplazada en {location}")
                            return True
            
            return False
            
            # # 1. Analizar el párrafo y crear un mapa de contenido
            # content_map = analyze_paragraph_content(paragraph)
            
            # # 2. Verificar si hay imágenes para reemplazar
            # images_to_replace = []
            
            # for item in content_map:
            #     if item['type'] == 'image':
            #         # Usar el primer reemplazo disponible de la configuración original (SIN agotar cola)
            #         placeholder_name, replacement_info = list(placeholder_replacements.items())[0]
            #         replacement_path = get_replacement_path(replacement_info)
                    
            #         print(f"🔍 Intentando reemplazar con: {replacement_path}")
                    
            #         if replacement_path and os.path.exists(replacement_path):
            #             print(f"✅ Archivo encontrado: {replacement_path}")
            #             item['replacement'] = {
            #                 'path': replacement_path,
            #                 'dimensions': calculate_replacement_dimensions(
            #                     replacement_path, replacement_info, item['original_dimensions']
            #                 )
            #             }
            #             images_to_replace.append(placeholder_name)
            #             break  # Solo reemplazar una imagen por párrafo
            #         else:
            #             print(f"❌ Archivo no existe: {replacement_path}")
            
            # if not images_to_replace:
            #     print(f"⚠️ No se pudo usar ningún reemplazo en {location}")
            #     return False
                
            # # 3. Reconstruir el párrafo con las nuevas imágenes
            # print(f"🔄 Reconstruyendo párrafo en {location}")
            # try:
            #     reconstruct_paragraph_with_replacements(paragraph, content_map)
            #     print(f"✅ Párrafo reconstruido exitosamente en {location}")
            #     for img_name in images_to_replace:
            #         print(f"   📸 Reemplazado: {img_name}")
            #     return True
            # except Exception as reconstruction_error:
            #     print(f"❌ Error durante reconstrucción: {reconstruction_error}")
            #     return False
            
        except Exception as e:
            print(f"❌ Error general procesando párrafo en {location}: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def analyze_paragraph_content(paragraph):
        """Analiza el contenido del párrafo y crea un mapa de elementos."""
        content_map = []
        
        for run in paragraph.runs:
            run_content = {'type': 'text', 'text': run.text, 'format': extract_run_format(run)}
            
            # Buscar imágenes en el run
            has_image = False
            for elem in run._element:
                if (elem.tag.endswith('drawing') or 
                    elem.tag.endswith('pict')):
                    # Este run contiene una imagen
                    image_info = extract_image_info(elem)
                    content_map.append({
                        'type': 'image',
                        'original_dimensions': image_info,
                        'element': elem
                    })
                    has_image = True
                    break
            
            # Solo añadir el texto si no hay imagen o si hay texto además de la imagen
            if not has_image and run.text:
                content_map.append(run_content)
            elif has_image and run.text:
                # Texto junto con imagen - dividir si es necesario
                content_map.append(run_content)
        
        return content_map
    
    def extract_run_format(run):
        """Extrae el formato de un run."""
        return {
            'font_name': run.font.name,
            'font_size': run.font.size,
            'bold': run.font.bold,
            'italic': run.font.italic,
            'underline': run.font.underline,
            'color': run.font.color.rgb if run.font.color.rgb else None
        }
    
    def reconstruct_paragraph_with_replacements(paragraph, content_map):
        """Reconstruye el párrafo con las imágenes reemplazadas."""
        # Guardar propiedades del párrafo original
        original_style = paragraph.style
        original_alignment = paragraph.alignment
        
        # Limpiar el párrafo
        for run in paragraph.runs[:]:
            run._element.getparent().remove(run._element)
        
        # Reconstruir contenido
        for item in content_map:
            if item['type'] == 'text' and item['text']:
                run = paragraph.add_run(item['text'])
                apply_format_to_run(run, item['format'])
                
            elif item['type'] == 'image':
                if 'replacement' in item:
                    # Insertar imagen de reemplazo
                    run = paragraph.add_run()
                    width = Cm(item['replacement']['dimensions']['width_cm'])
                    height = Cm(item['replacement']['dimensions']['height_cm'])
                    run.add_picture(item['replacement']['path'], width=width, height=height)
                else:
                    # Mantener imagen original (esto no debería pasar)
                    print("⚠️ Imagen original mantenida")
        
        # Restaurar propiedades del párrafo
        paragraph.style = original_style
        paragraph.alignment = original_alignment
    
    def extract_image_info(image_elem):
        """Extrae información de dimensiones de la imagen."""
        try:
            width_cm = 3.0
            height_cm = 2.0
            
            for child in image_elem.iter():
                if 'cx' in child.attrib and 'cy' in child.attrib:
                    width_cm = max(0.5, int(child.attrib['cx']) / 914400 * 2.54)
                    height_cm = max(0.5, int(child.attrib['cy']) / 914400 * 2.54)
                    break
            
            return {'width_cm': width_cm, 'height_cm': height_cm}
        except:
            return {'width_cm': 3.0, 'height_cm': 2.0}
    
    def calculate_replacement_dimensions(image_path, replacement_info, original_info):
        """Calcula dimensiones para la imagen de reemplazo."""
        if isinstance(replacement_info, dict):
            if 'width_cm' in replacement_info and 'height_cm' in replacement_info:
                return {
                    'width_cm': replacement_info['width_cm'],
                    'height_cm': replacement_info['height_cm']
                }
            
            if replacement_info.get('maintain_aspect', False):
                return calculate_aspect_ratio_dimensions(
                    image_path, original_info['width_cm'], original_info['height_cm']
                )
        
        return {
            'width_cm': original_info['width_cm'],
            'height_cm': original_info['height_cm']
        }
    
    def calculate_aspect_ratio_dimensions(image_path, max_width, max_height):
        """Calcula dimensiones manteniendo aspect ratio."""
        try:
            with Image.open(image_path) as img:
                img_width, img_height = img.size
                aspect_ratio = img_width / img_height
                
                if aspect_ratio > 1:
                    width = max_width
                    height = max_width / aspect_ratio
                    if height > max_height:
                        height = max_height
                        width = max_height * aspect_ratio
                else:
                    height = max_height
                    width = max_height * aspect_ratio
                    if width > max_width:
                        width = max_width
                        height = max_width / aspect_ratio
                
                return {'width_cm': width, 'height_cm': height}
        except:
            return {'width_cm': max_width, 'height_cm': max_height}
    
    def debug_document_structure():
        """Función de debugging para verificar la estructura del documento."""
        print("\n🔍 ANÁLISIS DETALLADO DEL DOCUMENTO:")
        print(f"   📄 Párrafos principales: {len(doc.paragraphs)}")
        print(f"   📊 Tablas principales: {len(doc.tables)}")
        print(f"   📋 Secciones: {len(doc.sections)}")
        
        # Verificar cada sección
        for i, section in enumerate(doc.sections):
            print(f"\n   📋 SECCIÓN {i+1}:")
            
            # Verificar headers
            headers = [
                (section.header, "header principal"),
                (section.first_page_header, "header primera página"),
                (section.even_page_header, "header páginas pares")
            ]
            
            for header, name in headers:
                if header is not None:
                    try:
                        has_paragraphs = (hasattr(header, 'paragraphs') and len(header.paragraphs) > 0)
                        has_tables = (hasattr(header, 'tables') and len(header.tables) > 0)
                        
                        print(f"      {name}:")
                        print(f"         📄 Párrafos: {len(header.paragraphs) if has_paragraphs else 0}")
                        print(f"         📊 Tablas: {len(header.tables) if has_tables else 0}")
                        
                        # Revisar párrafos
                        if has_paragraphs:
                            for p_idx, p in enumerate(header.paragraphs):
                                has_img = paragraph_has_images(p)
                                if has_img or p.text.strip():
                                    img_marker = "📷" if has_img else "📝"
                                    print(f"            {img_marker} Párrafo {p_idx+1}: '{p.text[:30]}...'")
                        
                        # Revisar tablas (ESTO ES CLAVE PARA TU CASO)
                        if has_tables:
                            for t_idx, table in enumerate(header.tables):
                                print(f"         📊 TABLA {t_idx+1}: {len(table.rows)} filas x {len(table.columns)} columnas")
                                
                                for r_idx, row in enumerate(table.rows):
                                    for c_idx, cell in enumerate(row.cells):
                                        cell_has_images = any(paragraph_has_images(p) for p in cell.paragraphs)
                                        cell_text = ' '.join(p.text for p in cell.paragraphs).strip()
                                        
                                        if cell_has_images or cell_text:
                                            img_marker = "📷" if cell_has_images else "📝"
                                            print(f"            {img_marker} Fila {r_idx+1}, Col {c_idx+1}: '{cell_text[:20]}...'")
                                            
                                            if cell_has_images:
                                                print(f"               ⚠️ ¡IMAGEN ENCONTRADA AQUÍ!")
                        
                    except Exception as e:
                        print(f"      {name}: ❌ Error: {e}")
                else:
                    print(f"      {name}: ❌ No existe")
        
        print("\n")
    
    def force_process_header_tables():
        """Función específica para procesar agresivamente las tablas en headers."""
        nonlocal total_replaced
        
        print("🔍 PROCESAMIENTO ESPECÍFICO DE TABLAS EN HEADERS...")
        
        # Crear una cola fresca específica para headers
        header_queue = list(placeholder_replacements.items())
        
        if not header_queue:
            print("⚠️ No hay reemplazos disponibles para headers")
            return
        
        print(f"📋 Reemplazos disponibles para headers: {len(header_queue)}")
        headers_processed = 0
        
        for i, section in enumerate(doc.sections):
            headers = [
                (section.header, "header principal"),
                (section.first_page_header, "header primera página"),
                (section.even_page_header, "header páginas pares")
            ]
            
            for header, name in headers:
                if header is not None and hasattr(header, 'tables') and header.tables:
                    section_has_images = False
                    
                    # Primero verificar si hay imágenes en esta sección
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    if paragraph_has_images(paragraph):
                                        section_has_images = True
                                        break
                                if section_has_images:
                                    break
                            if section_has_images:
                                break
                        if section_has_images:
                            break
                    
                    if section_has_images:
                        print(f"🎯 Procesando FORZADAMENTE tablas en {name} sección {i+1}")
                        headers_processed += 1
                        
                        for table_idx, table in enumerate(header.tables):
                            print(f"   📊 Tabla {table_idx+1}: {len(table.rows)}x{len(table.columns)}")
                            
                            for row_idx, row in enumerate(table.rows):
                                for col_idx, cell in enumerate(row.cells):
                                    for para_idx, paragraph in enumerate(cell.paragraphs):
                                        if paragraph_has_images(paragraph):
                                            location = f"{name} sección {i+1}, tabla {table_idx+1}, celda ({row_idx+1},{col_idx+1}), párrafo {para_idx+1}"
                                            print(f"      🎯 IMAGEN DETECTADA en {location}")
                                            
                                            if replace_image_direct_method(paragraph, location, header_queue):
                                                print(f"      ✅ IMAGEN REEMPLAZADA en {location}")
                                                total_replaced += 1
                                            else:
                                                print(f"      ❌ FALLÓ el reemplazo en {location}")
        
        print(f"📋 Headers procesados: {headers_processed}")

    def replace_image_direct_method(paragraph, location, available_replacements):
        """Método directo más simple para reemplazar imágenes."""
        try:
            if not available_replacements:
                print(f"      ⚠️ No hay reemplazos disponibles")
                return False
            
            # Tomar el primer reemplazo disponible (SIN REMOVERLO de la cola)
            placeholder_name, replacement_info = available_replacements[0]
            replacement_path = get_replacement_path(replacement_info)
            
            print(f"      🔄 Intentando reemplazar con: {os.path.basename(replacement_path) if replacement_path else 'N/A'}")
            
            if not replacement_path or not os.path.exists(replacement_path):
                print(f"      ❌ Archivo no encontrado: {replacement_path}")
                return False
            
            # Buscar la imagen en el párrafo y reemplazarla directamente
            for run_idx, run in enumerate(paragraph.runs):
                image_found = False
                for elem in run._element[:]:  # Crear copia para iterar
                    if (elem.tag.endswith('drawing') or 
                        elem.tag.endswith('pict') or 
                        'pic:pic' in str(elem.tag)):
                        
                        print(f"      🎯 Imagen encontrada en run {run_idx+1}, reemplazando...")
                        
                        # Extraer dimensiones originales
                        original_dims = extract_image_info_simple(elem)
                        
                        # Calcular nuevas dimensiones
                        new_dims = calculate_replacement_dimensions(
                            replacement_path, replacement_info, original_dims
                        )
                        
                        print(f"      📏 Dimensiones: {new_dims['width_cm']:.1f}x{new_dims['height_cm']:.1f}cm")
                        
                        # Método más directo: limpiar run y añadir imagen
                        try:
                            # Eliminar el elemento de imagen
                            elem.getparent().remove(elem)
                            
                            # Limpiar cualquier texto del run
                            run.text = ""
                            
                            # Añadir nueva imagen
                            width = Cm(new_dims['width_cm'])
                            height = Cm(new_dims['height_cm'])
                            run.add_picture(replacement_path, width=width, height=height)
                            
                            print(f"      ✅ Imagen insertada exitosamente usando {placeholder_name}")
                            image_found = True
                            return True
                            
                        except Exception as insert_error:
                            print(f"      ❌ Error insertando imagen: {insert_error}")
                            return False
                
                if image_found:
                    break
            
            print(f"      ⚠️ No se encontraron elementos de imagen válidos")
            return False
            
        except Exception as e:
            print(f"      ❌ Error en método directo: {e}")
            import traceback
            traceback.print_exc()
            return False

    def extract_image_info_simple(image_elem):
        """Extrae información de imagen de manera más robusta."""
        try:
            width_cm = 3.0
            height_cm = 2.0
            
            # Buscar dimensiones en diferentes niveles del XML
            for descendant in image_elem.iter():
                if hasattr(descendant, 'attrib'):
                    attrib = descendant.attrib
                    if 'cx' in attrib and 'cy' in attrib:
                        try:
                            width_cm = max(0.5, int(attrib['cx']) / 914400 * 2.54)
                            height_cm = max(0.5, int(attrib['cy']) / 914400 * 2.54)
                            break
                        except (ValueError, TypeError):
                            continue
            
            print(f"      📐 Dimensiones extraídas: {width_cm:.1f}x{height_cm:.1f}cm")
            return {'width_cm': width_cm, 'height_cm': height_cm}
            
        except Exception as e:
            print(f"      ⚠️ Error extrayendo dimensiones: {e}")
            return {'width_cm': 3.0, 'height_cm': 2.0}
    
    # Ejecutar debugging primero
    debug_document_structure()
    
    # Ejecutar procesamiento normal PRIMERO (para párrafos principales)
    print("📄 FASE 2A: Procesamiento normal del documento...")
    process_all_paragraphs()
    
    # Ejecutar procesamiento forzado de headers DESPUÉS (con cola fresca)
    print("📄 FASE 2B: Procesamiento específico de headers...")
    force_process_header_tables()
    
    print(f"🎯 Total de placeholders reemplazados: {total_replaced}")
    return total_replaced

def process_word_file(input_path: str, output_path: str, replacements: dict, 
                     image_replacements: dict = None, placeholder_replacements: dict = None):
    """Procesa archivo Word con sistema robusto de reemplazo de placeholders."""
    if image_replacements is None:
        image_replacements = {}
    if placeholder_replacements is None:
        placeholder_replacements = {}
    
    try:
        print(f"\n📄 ========== PROCESANDO: {os.path.basename(input_path)} ==========")
        doc = Document(input_path)
        
        if replacements:
            print("\n📝 FASE 1: Reemplazos de texto...")
            replace_text(doc, replacements, image_replacements)
            print("✅ Texto procesado")
        
        if placeholder_replacements:
            print("\n🖼️ FASE 2: Reemplazos de placeholders...")
            replaced_count = replace_placeholder_images(doc, placeholder_replacements.copy())
            print(f"✅ {replaced_count} placeholders procesados")
        
        print(f"\n💾 Guardando: {os.path.basename(output_path)}")
        doc.save(output_path)
        print("✅ ¡Documento guardado exitosamente!")
        print("=" * 60 + "\n")
        
    except Exception as e:
        print(f"❌ ERROR: {e}")
        raise

def create_placeholder_images(output_dir="assets"):
    """Crea imágenes placeholder."""
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    placeholders = [
        {"name": "placeholder_logo.png", "text": "LOGO", "size": (400, 200)},
        {"name": "placeholder_firma.png", "text": "FIRMA", "size": (300, 100)},
        {"name": "placeholder_foto.png", "text": "FOTO", "size": (200, 250)}
    ]
    
    for placeholder in placeholders:
        try:
            from PIL import Image, ImageDraw, ImageFont
            
            img = Image.new('RGB', placeholder['size'], color='#f0f0f0')
            draw = ImageDraw.Draw(img)
            
            try:
                font = ImageFont.truetype("arial.ttf", 36)
            except:
                font = ImageFont.load_default()
            
            bbox = draw.textbbox((0, 0), placeholder['text'], font=font)
            x = (placeholder['size'][0] - (bbox[2] - bbox[0])) // 2
            y = (placeholder['size'][1] - (bbox[3] - bbox[1])) // 2
            
            draw.text((x, y), placeholder['text'], fill='#888888', font=font)
            draw.rectangle([0, 0, placeholder['size'][0]-1, placeholder['size'][1]-1], 
                          outline='#cccccc', width=2)
            
            filepath = os.path.join(output_dir, placeholder['name'])
            img.save(filepath)
            print(f"✅ Placeholder: {filepath}")
            
        except Exception as e:
            print(f"❌ Error: {placeholder['name']}: {e}")

if __name__ == "__main__":
    create_placeholder_images()
    
    replacements = {"{{NOMBRE}}": "Juan Pérez"}
    placeholder_replacements = {
        "placeholder_logo.png": {
            "path": "mi_logo.png",
            "maintain_aspect": True
        }
    }
    
    process_word_file("template.docx", "output.docx", replacements, {}, placeholder_replacements)